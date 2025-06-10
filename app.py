import streamlit as st
import pandas as pd
import io
import zipfile
from collections import defaultdict
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openpyxl import Workbook
from openpyxl.styles import PatternFill

st.set_page_config(page_title="SeatMatrix GGSIPU", layout="wide")
st.title("📘 SeatMatrix GGSIPU - Exam Seating Planner")

# File uploads
excel_file = st.file_uploader("📄 Upload Student Excel", type=["xlsx"])
template_docx = st.file_uploader("📄 Upload Word Template", type=["docx"])

# Input fields
mapping_input = st.text_input("🔗 Enter Mappings", "ICT202-16407722-16407255-ECE")
room_input = st.text_input("🏫 Enter Room Specs", "Room1:48:6x8")
date_input = st.text_input("📅 Enter Date", "31-05-2025")
time_input = st.text_input("⏰ Enter Time", "10:00 AM – 1:00 PM")

if st.button("🚀 Generate Seating Plan") and excel_file and template_docx:
    df = pd.read_excel(excel_file)
    df.columns = df.columns.str.strip().str.lower()
    df = df[['name', 'rollno', 'paper code']]
    df['rollno'] = df['rollno'].astype(str).str.zfill(11)
    df['paper code'] = df['paper code'].str.strip()
    df['last8'] = df['rollno'].str[-8:]

    paper_last8_dept_map = {}
    for entry in mapping_input.split(","):
        parts = entry.strip().split("-")
        if len(parts) >= 3:
            paper, dept = parts[0].strip(), parts[-1].strip()
            for last8 in parts[1:-1]:
                paper_last8_dept_map[(paper, last8.strip())] = dept

    valid_papers = {k[0] for k in paper_last8_dept_map}
    df = df[df['paper code'].isin(valid_papers)]
    df['department'] = df.apply(lambda row: paper_last8_dept_map.get((row['paper code'], row['last8'])), axis=1)
    df = df[df['department'].notna()]

    room_specs = room_input.split(",")
    parsed_rooms = []
    for spec in room_specs:
        parts = spec.strip().split(":")
        name = parts[0]
        layout = parts[2] if len(parts) == 3 else "6x8"
        rows, cols = map(int, layout.lower().split("x"))
        parsed_rooms.append((name, rows, cols))

    paper_groups = defaultdict(list)
    for _, row in df.iterrows():
        paper_groups[row['paper code']].append((row['rollno'], row['department']))

    color_palette = ["F8CBAD", "DDEBF7", "C6E0B4", "F4B084", "FFD966", "D9D2E9", "B4C6E7", "E2EFDA"]
    paper_colors = {paper: color_palette[i % len(color_palette)] for i, paper in enumerate(paper_groups)}

    def fill_columnwise(paper_queue, paper_groups, rows, cols):
        room = [["" for _ in range(cols)] for _ in range(rows)]
        dept_map = [["" for _ in range(cols)] for _ in range(rows)]
        paper_map = [["" for _ in range(cols)] for _ in range(rows)]
        seat_order = [(r, c) for c in range(cols) for r in range(rows)]
        seat_index = 0
        while seat_index < len(seat_order) and paper_queue:
            p1 = paper_queue[0]
            p2 = paper_queue[1] if len(paper_queue) > 1 else None
            for i in range(seat_index, len(seat_order)):
                r, c = seat_order[i]
                current_paper = None
                if c % 2 == 0 and paper_groups[p1]:
                    current_paper = p1
                elif c % 2 == 1 and p2 and paper_groups[p2]:
                    current_paper = p2
                if current_paper:
                    roll, dept = paper_groups[current_paper].pop(0)
                    room[r][c] = roll
                    dept_map[r][c] = dept
                    paper_map[r][c] = current_paper
                    seat_index += 1
                    if not paper_groups[current_paper]:
                        paper_queue.remove(current_paper)
                    break
                else:
                    seat_index += 1
        return room, dept_map, paper_map

    def get_column_departments(dept_map, col, rows):
        return "/".join(sorted({dept_map[r][col].upper() for r in range(rows) if dept_map[r][col]}))

    def set_table_borders(table):
        for row in table.rows:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                borders = OxmlElement('w:tcBorders')
                for edge in ('top', 'left', 'bottom', 'right'):
                    tag = OxmlElement(f'w:{edge}')
                    tag.set(qn('w:val'), 'single')
                    tag.set(qn('w:sz'), '4')
                    tag.set(qn('w:space'), '0')
                    tag.set(qn('w:color'), '000000')
                    borders.append(tag)
                tcPr.append(borders)

    wb = Workbook()
    wb.remove(wb.active)
    paper_queue = list(paper_groups.keys())
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w') as zipf:
        for room_name, rows, cols in parsed_rooms:
            if not any(paper_groups.values()):
                break
            room, dept_map, paper_map = fill_columnwise(paper_queue, paper_groups, rows, cols)
            template_bytes = io.BytesIO(template_docx.read())
            doc = Document(template_bytes)

            for p in doc.paragraphs:
                if 'DATE:' in p.text:
                    p.text = f'DATE: {date_input}'
                elif 'TIME' in p.text.upper():
                    p.text = f'TIME: {time_input}'
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    if p.runs:
                        p.runs[0].bold = True
                elif 'ROOM NO.' in p.text.upper():
                    p.clear()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run1 = p.add_run("SEATING ARRANGEMENT FOR ")
                    run2 = p.add_run(f"Room No. {room_name}")
                    run2.bold = True
                    run2.font.size = Pt(14)

            summary_count = defaultdict(int)
            for r in range(rows):
                for c in range(cols):
                    roll = room[r][c]
                    dept = dept_map[r][c]
                    paper = paper_map[r][c]
                    if roll and dept and paper:
                        summary_count[(dept, paper)] += 1

            for para in doc.paragraphs:
                if "PAPER CODE" in para.text:
                    para.clear()
                    break

            for (dept, paper), count in summary_count.items():
                para = doc.add_paragraph(f"{dept.upper()} (PAPER CODE {paper}) – {{{count}}}")
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.runs[0].bold = True

            table = doc.add_table(rows=rows + 1, cols=cols)
            table.style = 'Table Grid'
            for c in range(cols):
                dept = get_column_departments(dept_map, c, rows)
                col_label = "ROW-1" if c < cols // 2 else "ROW-2"
                table.cell(0, c).text = f"{dept}\n{col_label}"
                para = table.cell(0, c).paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.runs[0].bold = True
            for r in range(rows):
                for c in range(cols):
                    table.cell(r + 1, c).text = room[r][c]
                    table.cell(r + 1, c).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_table_borders(table)

            word_name = f"{room_name}_Seating.docx"
            word_buffer = io.BytesIO()
            doc.save(word_buffer)
            zipf.writestr(word_name, word_buffer.getvalue())

            sheet = wb.create_sheet(title=room_name)
            for c in range(cols):
                dept = get_column_departments(dept_map, c, rows)
                col_label = "ROW-1" if c < cols // 2 else "ROW-2"
                sheet.cell(row=1, column=c + 2, value=f"{dept}\n{col_label}")
            for r in range(rows):
                sheet.cell(row=r + 2, column=1, value=f"Row {r+1}")
                for c in range(cols):
                    roll = room[r][c]
                    paper = paper_map[r][c]
                    cell = sheet.cell(row=r + 2, column=c + 2, value=roll)
                    if paper:
                        fill_color = paper_colors.get(paper, "FFFFFF")
                        cell.fill = PatternFill(start_color=fill_color, end_color=fill_color, fill_type="solid")

        excel_buffer = io.BytesIO()
        wb.save(excel_buffer)
        zipf.writestr("Seating_Summary.xlsx", excel_buffer.getvalue())

    st.success("✅ Seating plan generated successfully!")
    st.download_button("⬇️ Download Seating Plan ZIP", data=zip_buffer.getvalue(), file_name="Final_Seating_Documents.zip")
